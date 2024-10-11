import curses
import os
import textwrap
from docx import Document as DocxDocument
import ollama

def list_documents():
    try:
        return [f for f in os.listdir('documents') if os.path.isfile(os.path.join('documents', f))]
    except FileNotFoundError:
        return []

def read_file(file_path):
    try:
        if file_path.endswith('.docx'):
            doc = DocxDocument(file_path)
            return '\n\n'.join([para.text for para in doc.paragraphs])
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
    except (UnicodeDecodeError, FileNotFoundError) as e:
        print(f"Error reading {file_path}: {e}")
        return ""
    except Exception as e:
        print(f"Unexpected error: {e}")
        return ""

def write_file(file_path, paragraphs):
    if file_path.endswith('.docx'):
        doc = DocxDocument()
        for para in paragraphs:
            doc.add_paragraph(para)
        doc.save(file_path)
    else:
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n\n'.join(paragraphs))

def segment_paragraphs(text):
    """
    Segments the text by line breaks to detect each new line as a segment.
    """
    # Split the text by single line breaks to handle individual lines as segments
    paragraphs = text.splitlines()
    
    # Remove empty lines if any
    paragraphs = [para for para in paragraphs if para.strip() != '']
    
    return paragraphs


def translate_text_with_ollama(text, source_language, target_language, model_name):
    prompt = f"Only provide the translation and nothing else. Translate this {source_language} text into {target_language}: {text}"
    response = ollama.chat(
        model=model_name,
        messages=[{'role': 'user', 'content': prompt}],
        stream=False
    )
    return response['message']['content']

def get_available_models():
    models_response = ollama.list()
    # Filter out models with 'embed' in their names
    models = [model['name'] for model in models_response['models'] if 'embed' not in model['name'].lower()]
    return models

def wrap_text(text, width):
    return textwrap.wrap(text, width)

def load_languages():
    try:
        with open('languages.txt', 'r') as f:
            languages = [line.strip() for line in f if line.strip()]
        return languages
    except FileNotFoundError:
        print("languages.txt not found.")
        return []

def display_translation_tool(screen):
    curses.curs_set(0)  # Hide the cursor
    height, width = screen.getmaxyx()

    if height < 20 or width < 80:
        screen.addstr(0, 0, "Terminal window too small. Please resize.")
        screen.refresh()
        screen.getch()
        return

    # Initialize color pairs
    curses.start_color()
    curses.init_pair(1, curses.COLOR_RED, curses.COLOR_BLACK)  # Color pair 1 for red text

    left_win = screen.subwin(height - 6, width // 2, 0, 0)
    right_win = screen.subwin(height - 6, width // 2, 0, width // 2)
    status_win = screen.subwin(2, width, height - 2, 0)

    languages = load_languages()
    if not languages:
        screen.addstr(0, 0, "No languages found in languages.txt.")
        screen.refresh()
        screen.getch()
        return

    models = get_available_models()

    source_lang = 0
    target_lang = 0
    selected_model = 0
    selecting = 'language'
    selected_doc = 0
    model_name = None

    def refresh_status():
        status_win.clear()
        status_win.addstr(1, 0, "Press 'q' to quit.")
        status_win.refresh()

    def display_language_selection():
        screen.clear()
        screen.addstr(0, 0, "Select source language:")
        for idx, lang in enumerate(languages):
            if idx == source_lang:
                screen.addstr(idx + 2, 2, f"> {lang}", curses.A_BOLD)
            else:
                screen.addstr(idx + 2, 2, lang)

        screen.addstr(len(languages) + 3, 0, "Select target language:")
        for idx, lang in enumerate(languages):
            if idx == target_lang:
                screen.addstr(len(languages) + 5 + idx, 2, f"> {lang}", curses.A_BOLD)
            else:
                screen.addstr(len(languages) + 5 + idx, 2, lang)

    def display_document_selection():
        documents = list_documents()
        screen.clear()
        screen.addstr(0, 0, "Select a document to translate:")
        for idx, doc in enumerate(documents):
            if idx == selected_doc:
                screen.addstr(idx + 2, 2, f"> {doc}", curses.A_BOLD)
            else:
                screen.addstr(idx + 2, 2, doc)

    def display_model_selection():
        global models, selected_model  # Make sure models is a global variable
        models = get_available_models()  # Refresh models list
        if not models:
            selected_model = None  # Handle empty model list
        current_model = 0 if models else None  # Only set a model if there are models
        
        while True:
            screen.clear()
            if models:
                for idx, model in enumerate(models):
                    if idx == current_model:
                        screen.addstr(idx + 1, 0, f"> {model}", curses.A_BOLD)
                    else:
                        screen.addstr(idx + 1, 0, f"  {model}")
                screen.addstr(len(models) + 2, 0, "Press 'Enter' to select a model, 'r' to refresh")
            else:
                screen.addstr(0, 0, "No models available.")
                screen.addstr(1, 0, "Press 'r' to refresh")

            screen.refresh()

            key = screen.getch()
            if key == curses.KEY_UP:
                if models:
                    current_model = (current_model - 1) % len(models)
            elif key == curses.KEY_DOWN:
                if models:
                    current_model = (current_model + 1) % len(models)
            elif key == 10:  # Enter key to select a model
                if models and current_model is not None:
                    global model_name
                    model_name = models[current_model]
                    return
            elif key == ord('r'):  # 'r' key to refresh model list
                display_model_selection()  # Refresh model list

    def process_key(key):
        nonlocal selecting, source_lang, target_lang, selected_doc, selected_model, model_name

        if selecting == 'language':
            if key == curses.KEY_UP:
                source_lang = (source_lang - 1) % len(languages)
            elif key == curses.KEY_DOWN:
                source_lang = (source_lang + 1) % len(languages)
            elif key == curses.KEY_LEFT:
                target_lang = (target_lang - 1) % len(languages)
            elif key == curses.KEY_RIGHT:
                target_lang = (target_lang + 1) % len(languages)
            elif key == 10:  # Enter key to confirm language selection
                selecting = 'document'
        elif selecting == 'document':
            documents = list_documents()
            if key == curses.KEY_UP:
                selected_doc = (selected_doc - 1) % len(documents) if documents else 0
            elif key == curses.KEY_DOWN:
                selected_doc = (selected_doc + 1) % len(documents) if documents else 0
            elif key == 10:  # Enter key to select a document
                selecting = 'model'
        elif selecting == 'model':
            if key == curses.KEY_UP:
                if models:
                    selected_model = (selected_model - 1) % len(models) if models else 0
            elif key == curses.KEY_DOWN:
                if models:
                    selected_model = (selected_model + 1) % len(models) if models else 0
            elif key == 10:  # Enter key to confirm model selection
                if models and selected_model is not None:
                    model_name = models[selected_model]
                    selecting = 'translation'  # Move to translation stage
            elif key == ord('r'):  # 'r' key to refresh model list
                display_model_selection()

    def start_translation():
        nonlocal model_name, source_lang, target_lang
        source_language = languages[source_lang]
        target_language = languages[target_lang]

        documents = list_documents()
        file_path = os.path.join('documents', documents[selected_doc])
        content = read_file(file_path)

        paragraphs = segment_paragraphs(content)
        current_paragraph = 0
        translations = ['' for _ in paragraphs]
        scroll_position = 0  # To control the scroll of paragraphs

        translation_dir = 'translations'
        os.makedirs(translation_dir, exist_ok=True)
        file_ext = os.path.splitext(documents[selected_doc])[1]
        translation_file_path = os.path.join(translation_dir, f"{os.path.splitext(documents[selected_doc])[0]}_{target_language}{file_ext}")

        if os.path.exists(translation_file_path):
            translated_content = read_file(translation_file_path).split('\n\n')
            for i, translation in enumerate(translated_content):
                if i < len(translations):
                    translations[i] = translation

        while True:
            left_win.clear()
            right_win.clear()

            gap = "    "  # Define a gap between the windows
            display_height = height - 6  # Height available for displaying lines

            # Wrap the paragraphs and translations for the left and right windows
            wrapped_paragraphs = [wrap_text(para, width // 2 - 4) for para in paragraphs]
            wrapped_translations = [wrap_text(translations[i], width // 2 - 4) for i in range(len(translations))]

            # Create a list of lines to display for both the source (left) and translations (right)
            left_display_lines = []
            right_display_lines = []

            for i, wrapped_paragraph in enumerate(wrapped_paragraphs):
                if i == current_paragraph:
                    # Add `>` indicator for the current paragraph
                    left_display_lines.append([f"> {wrapped_paragraph[0]}"] + wrapped_paragraph[1:])
                else:
                    left_display_lines.append(wrapped_paragraph)

                # Add corresponding translation or blank if no translation
                right_display_lines.append(wrapped_translations[i] if i < len(wrapped_translations) else [''])

                # Add a blank line to separate paragraphs visually
                left_display_lines.append([""])
                right_display_lines.append([""])

            # Flatten the lists
            left_display_lines = [line for para in left_display_lines for line in para]
            right_display_lines = [line for para in right_display_lines for line in para]

            # Ensure scroll position is correct
            max_scroll = max(0, len(left_display_lines) - display_height)
            scroll_position = min(scroll_position, max_scroll)
            scroll_position = max(scroll_position, 0)

            # Display paragraphs and translations
            for i in range(scroll_position, min(scroll_position + display_height, len(left_display_lines))):
                try:
                    left_line = left_display_lines[i] if i < len(left_display_lines) else ""
                    right_line = right_display_lines[i] if i < len(right_display_lines) else ""

                    # Display left window text
                    left_win.addstr(i - scroll_position, 0, left_line)

                    # Display right window text
                    right_win.addstr(i - scroll_position, 0, gap + right_line)
                except curses.error:
                    pass

            # Display word count of the current paragraph with conditional coloring
            current_paragraph_word_count = len(paragraphs[current_paragraph].split())
            status_win.clear()
            # Correctly format the paragraph count and word count
            paragraph_info = f"Paragraph {current_paragraph + 1}/{len(paragraphs)} - "
            word_count_info = f"Word count: {current_paragraph_word_count}"

            status_win.addstr(0, 0, paragraph_info)
            status_win.addstr(word_count_info, curses.color_pair(1) if current_paragraph_word_count > 1000 else curses.A_NORMAL)
            status_win.addstr(1, 0, "Press 'q' to quit.")
            status_win.refresh()

            left_win.refresh()
            right_win.refresh()

            key = screen.getch()

            # Scroll through paragraphs
            if key == curses.KEY_UP:
                current_paragraph = max(0, current_paragraph - 1)  # Move to previous paragraph
            elif key == curses.KEY_DOWN:
                current_paragraph = min(len(paragraphs) - 1, current_paragraph + 1)  # Move to next paragraph
            elif key == 10:  # Enter key to translate the selected paragraph
                if paragraphs[current_paragraph].strip():
                    translated = translate_text_with_ollama(paragraphs[current_paragraph], source_language, target_language, model_name)
                    translations[current_paragraph] = translated
                else:
                    translations[current_paragraph] = ''

            elif key == curses.KEY_RIGHT:  # Right arrow key to send text as is without translation
                translations[current_paragraph] = paragraphs[current_paragraph]

            # Update the scroll position based on the current paragraph
            scroll_position = sum(len(wrap_text(paragraphs[i], width // 2 - 4)) + 1 for i in range(current_paragraph))

            # Save translations to a file
            write_file(translation_file_path, translations)

            if key == ord('q'):
                return

    while True:
        if selecting == 'language':
            display_language_selection()
        elif selecting == 'document':
            display_document_selection()
        elif selecting == 'model':
            display_model_selection()
        elif selecting == 'translation':
            start_translation()

        refresh_status()
        key = screen.getch()
        if key == ord('q'):
            return
        process_key(key)

if __name__ == "__main__":
    curses.wrapper(display_translation_tool)
